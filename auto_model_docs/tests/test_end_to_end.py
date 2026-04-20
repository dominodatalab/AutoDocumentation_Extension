"""End-to-end pipeline, notebook round-trip, and sanitizer invariant tests.

These tests exercise the full pipeline with a fake LLM (no network calls) and
the real scanner, planner, generator, builder, and notebook modules.
"""

from __future__ import annotations

import asyncio
import io
from pathlib import Path
from typing import Any, Dict, List, Optional
from unittest.mock import patch

import nbformat
import pytest
from docx import Document as DocxDocument

import artifact_layout
import dataset_store

from autodoc.core.models import (
    ArtifactContext,
    DocumentSpec,
    SectionSpec,
)
from autodoc.generation.notebook_exporter import NotebookExporter
from autodoc.llm.client import LLMResponse
from autodoc.orchestrator import Orchestrator
from autodoc.scanning.sanitizer import ContentSanitizer


# ---------------------------------------------------------------------------
# In-memory DatasetStore (mirrors the pattern used in test_orchestrator.py)
# ---------------------------------------------------------------------------


class _MemStore:
    dataset_id = "ds-test"
    snapshot_id = "snap-test"

    def __init__(self) -> None:
        self._files: Dict[str, bytes] = {}

    def write_file(self, path: str, content: bytes) -> None:
        self._files[path] = content

    def read_file(self, path: str) -> bytes:
        if path not in self._files:
            raise FileNotFoundError(path)
        return self._files[path]

    def file_exists(self, path: str) -> bool:
        return path in self._files

    def list_files(self, path: str = "") -> list:
        return []


@pytest.fixture
def mem_store() -> _MemStore:
    artifact_layout.init_layout()
    store = _MemStore()
    dataset_store._store = store
    yield store
    artifact_layout.reset_layout()
    dataset_store.reset_store()


# ---------------------------------------------------------------------------
# Fake LLM
# ---------------------------------------------------------------------------


_RANKING = {
    "ranked_files": [
        {"path": "train.py", "role": "training", "confidence": 0.95},
        {"path": "features.py", "role": "preprocessing", "confidence": 0.8},
    ]
}

_CODE_ANALYSIS = {
    "model_classes": ["XGBClassifier"],
    "features": ["age", "income", "credit_score"],
    "target_variable": "default",
    "transformations": [],
    "ml_task_type": "classification",
    "hyperparameters": {"n_estimators": 100, "max_depth": 5},
    "data_sources": ["customers.csv"],
    "insights": "XGBoost classifier trained on customer features.",
    "code_evidence": [
        {
            "statement": "XGBClassifier is the model",
            "file": "train.py",
            "symbol": "train_model",
            "snippet": "model = XGBClassifier(n_estimators=100, max_depth=5)",
        }
    ],
}

_PLAN_NARRATIVE_ONLY = {
    "section_title": "Overview",
    "content_blocks": [
        {
            "type": "narrative",
            "purpose": "Summarize the model and its features",
            "data_needed": "code analysis",
            "specifics": {},
        }
    ],
}

_NARRATIVE_TEXT = (
    "The project trains an XGBClassifier on customer features to predict default risk. "
    "Hyperparameters are fixed at n_estimators=100 and max_depth=5."
)


class FakeLLM:
    """Fake LLM client implementing complete and complete_json.

    Dispatches based on schema shape. Optionally records every call so
    tests can assert on what was sent to the LLM.
    """

    def __init__(self, record: bool = False) -> None:
        self.record = record
        self.calls: List[Dict[str, str]] = []

    async def complete_json(
        self,
        prompt: str,
        schema: Dict[str, Any],
        system: str = "",
        **_: Any,
    ) -> Dict[str, Any]:
        if self.record:
            self.calls.append({"prompt": prompt, "system": system})

        props = schema.get("properties") or {}
        required = schema.get("required") or []

        if "ranked_files" in required:
            return _RANKING
        if "code_evidence" in props and "model_classes" in props:
            return _CODE_ANALYSIS
        if "section_title" in required and "content_blocks" in required:
            return _PLAN_NARRATIVE_ONLY
        if "rows" in props and "columns" in props:
            return {
                "caption": "Data",
                "columns": ["Feature", "Value"],
                "rows": [{"Feature": "age", "Value": 1}],
            }
        if "labels" in props and "values" in props:
            return {
                "title": "Chart",
                "labels": ["a", "b"],
                "values": [1.0, 2.0],
                "xlabel": "x",
                "ylabel": "y",
            }
        if "items" in props:
            return {"title": "Items", "items": ["one", "two", "three"]}
        return {}

    async def complete(
        self,
        prompt: str,
        system: str = "",
        max_tokens: int = 4096,
        temperature: float = 0.7,
    ) -> LLMResponse:
        if self.record:
            self.calls.append({"prompt": prompt, "system": system})
        return LLMResponse(
            content=_NARRATIVE_TEXT,
            input_tokens=10,
            output_tokens=20,
            model="fake-model",
        )


# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------


@pytest.fixture
def sample_code_root(tmp_path: Path) -> Path:
    (tmp_path / "train.py").write_text(
        '"""Training script."""\n'
        "import pandas as pd\n"
        "from xgboost import XGBClassifier\n\n"
        "def train_model(X, y):\n"
        "    model = XGBClassifier(n_estimators=100, max_depth=5)\n"
        "    model.fit(X, y)\n"
        "    return model\n"
    )
    (tmp_path / "features.py").write_text(
        '"""Feature engineering."""\n'
        "def make_features(df):\n"
        "    return df[['age', 'income', 'credit_score']]\n"
    )
    (tmp_path / "README.md").write_text("# Credit risk model\n")
    return tmp_path


@pytest.fixture
def empty_artifact_context():
    """Patch ArtifactScanner.scan to return an empty ArtifactContext (no MLflow)."""

    async def _empty_scan(self, on_progress=None):
        if on_progress:
            on_progress(1.0)
        return ArtifactContext(
            models=[],
            datasets=[],
            project_metadata={"mlflow_available": False},
        )

    with patch(
        "autodoc.scanning.artifact_scanner.ArtifactScanner.scan",
        _empty_scan,
    ):
        yield


@pytest.fixture
def simple_spec() -> DocumentSpec:
    return DocumentSpec(
        title="Test Model Documentation",
        authors="Test Suite",
        sections=[SectionSpec(name="Overview")],
    )


# ---------------------------------------------------------------------------
# Test 1: End-to-end smoke
# ---------------------------------------------------------------------------


def test_pipeline_smoke_end_to_end(
    mem_store: _MemStore,
    empty_artifact_context,
    sample_code_root: Path,
    simple_spec: DocumentSpec,
) -> None:
    """Full pipeline runs with a fake LLM and writes a readable .docx.

    Exercises: language detection, code scanning (ranking + analysis),
    planning, content generation (narrative), and Word document assembly.
    """
    llm = FakeLLM()
    orch = Orchestrator(
        llm=llm,
        sanitizer=ContentSanitizer(),
        code_root=sample_code_root,
        output_dir=Path("docs"),
        parallel_workers=1,
        planning_workers=1,
        scan_workers=1,
        batch_size=4,
    )

    output_path = asyncio.new_event_loop().run_until_complete(
        orch.generate(simple_spec)
    )

    assert output_path.endswith(".docx")
    assert mem_store.file_exists(output_path), (
        f"Expected docx at {output_path}, mem store has {list(mem_store._files)}"
    )

    doc_bytes = mem_store.read_file(output_path)
    doc = DocxDocument(io.BytesIO(doc_bytes))
    all_text = "\n".join(p.text for p in doc.paragraphs)

    assert "Test Model Documentation" in all_text
    assert "Overview" in all_text
    assert "XGBClassifier" in all_text or "default risk" in all_text


# ---------------------------------------------------------------------------
# Test 2: Notebook round-trip
# ---------------------------------------------------------------------------


def test_notebook_roundtrip(
    mem_store: _MemStore,
    empty_artifact_context,
    sample_code_root: Path,
    simple_spec: DocumentSpec,
) -> None:
    """Edit a notebook cell, re-export to Word, verify the edit appears.

    This verifies the product's core editability promise: users can modify
    the generated .ipynb and round-trip the changes back into a .docx.
    """
    llm = FakeLLM()

    def _skip_execute(self, nb):
        return nb

    with patch(
        "autodoc.generation.notebook_builder.NotebookBuilder._execute_notebook",
        _skip_execute,
    ):
        orch = Orchestrator(
            llm=llm,
            sanitizer=ContentSanitizer(),
            code_root=sample_code_root,
            output_dir=Path("docs"),
            generate_notebook=True,
            parallel_workers=1,
            planning_workers=1,
            scan_workers=1,
            batch_size=4,
        )
        asyncio.new_event_loop().run_until_complete(orch.generate(simple_spec))

    notebook_paths = [p for p in mem_store._files if p.endswith(".ipynb")]
    assert len(notebook_paths) == 1, (
        f"Expected exactly one notebook, found: {notebook_paths}"
    )
    notebook_path = notebook_paths[0]

    nb_bytes = mem_store.read_file(notebook_path)
    nb = nbformat.read(io.StringIO(nb_bytes.decode("utf-8")), as_version=4)

    marker = "EDITED_BY_TEST_MARKER_4242"
    edited_any = False
    saw_section_header = False
    for cell in nb.cells:
        if (
            cell.cell_type == "markdown"
            and cell.source.lstrip().startswith("## 1.")
        ):
            saw_section_header = True
            continue
        if not saw_section_header:
            continue
        if cell.cell_type != "markdown":
            continue
        if "Export to Word" in cell.source:
            break
        cell.source = cell.source + f"\n\n{marker}\n"
        edited_any = True
        break
    assert edited_any, "Could not find a content markdown cell to edit"

    buf = io.StringIO()
    nbformat.write(nb, buf)
    mem_store.write_file(notebook_path, buf.getvalue().encode("utf-8"))

    exporter = NotebookExporter(output_dir=Path("docs"))
    exporter.export_to_word(Path(notebook_path), title="Test Model Documentation")

    docx_paths = [
        p for p in mem_store._files
        if p.endswith(".docx") and p != notebook_path
    ]
    assert docx_paths, "No docx produced by notebook export"
    exported_docx = sorted(docx_paths)[-1]

    doc = DocxDocument(io.BytesIO(mem_store.read_file(exported_docx)))
    all_text = "\n".join(p.text for p in doc.paragraphs)
    assert marker in all_text, (
        f"Edit did not survive round-trip. Docx text: {all_text[:500]}"
    )


# ---------------------------------------------------------------------------
# Test 3: Sanitizer invariant (no secrets reach the LLM)
# ---------------------------------------------------------------------------


# Known secret values that MUST NEVER appear in any LLM prompt.
_SECRETS = {
    "anthropic_key": "sk-ant-api03-THISISAFAKEANTHROPICKEYVALUE1234567890",
    "openai_key": "sk-proj-DEADBEEFDEADBEEFDEADBEEFDEADBEEF1234",
    "aws_access_key": "AKIAIOSFODNN7EXAMPLE",
    "aws_secret": "wJalrXUtnFEMI/K7MDENG/bPxRfiCYEXAMPLEKEY",
    "slack_token": "xoxb-1234567890-1234567890123-abcdefghij0123456789",
    "github_token": "ghp_AbCdEfGhIjKlMnOpQrStUvWxYz0123456789",
    "db_password": "superSecretPassword99!",
    "db_url": "postgres://admin:hunter2@10.0.5.42:5432/production",
}


def test_sanitizer_invariant_no_secrets_in_llm_prompts(
    mem_store: _MemStore,
    empty_artifact_context,
    tmp_path: Path,
    simple_spec: DocumentSpec,
) -> None:
    """No known secret value should appear in anything sent to the LLM.

    Populates a code root with a variety of secrets across regular source
    files, config files, and sensitive file types (.env, .pem, credentials),
    then captures every LLM prompt and asserts none of the literal secrets
    appear verbatim in the prompt or system message.
    """
    (tmp_path / "train.py").write_text(
        '"""Training script."""\n'
        "import os\n\n"
        f'ANTHROPIC_API_KEY = "{_SECRETS["anthropic_key"]}"\n'
        f'AWS_ACCESS_KEY_ID = "{_SECRETS["aws_access_key"]}"\n'
        f'AWS_SECRET_ACCESS_KEY = "{_SECRETS["aws_secret"]}"\n'
        "def train():\n"
        "    pass\n"
    )
    (tmp_path / "features.py").write_text(
        '"""Features."""\n'
        f'SLACK_TOKEN = "{_SECRETS["slack_token"]}"\n'
        f'GITHUB_TOKEN = "{_SECRETS["github_token"]}"\n'
        "def extract(df):\n"
        "    return df\n"
    )
    (tmp_path / "config.py").write_text(
        f'DATABASE_URL = "{_SECRETS["db_url"]}"\n'
        f'db_password = "{_SECRETS["db_password"]}"\n'
    )
    (tmp_path / ".env").write_text(
        f"OPENAI_API_KEY={_SECRETS['openai_key']}\n"
        f"ANTHROPIC_API_KEY={_SECRETS['anthropic_key']}\n"
    )
    (tmp_path / "credentials.yaml").write_text(
        f"api_key: {_SECRETS['openai_key']}\n"
    )

    llm = FakeLLM(record=True)
    orch = Orchestrator(
        llm=llm,
        sanitizer=ContentSanitizer(),
        code_root=tmp_path,
        output_dir=Path("docs"),
        parallel_workers=1,
        planning_workers=1,
        scan_workers=1,
        batch_size=4,
    )

    asyncio.new_event_loop().run_until_complete(orch.generate(simple_spec))

    assert llm.calls, "Expected at least one LLM call during generation"

    leaks: List[str] = []
    for call in llm.calls:
        haystack = (call["prompt"] or "") + "\n" + (call["system"] or "")
        for name, secret in _SECRETS.items():
            if secret in haystack:
                leaks.append(f"{name} ({secret[:10]}...) leaked in an LLM call")

    assert not leaks, "Secrets reached the LLM:\n" + "\n".join(leaks)
